import pandas as pd
import os
import re
import PyPDF2
from docx import Document
import logging
import tempfile
import shutil
from zipfile import ZipFile, BadZipFile
import csv
import openpyxl
import xlrd
import pdfplumber
from PIL import Image
import base64
import io
from dotenv import load_dotenv
import requests

load_dotenv()



# Налаштування логування
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def convert_to_csv(input_path, output_path=None):
    """
    Конвертує файли різних форматів (.docx, .doc, .xls, .xlsx, pdf) у CSV формат.
    
    Args:
        input_path: Шлях до вхідного файлу
        output_path: Шлях до вихідного CSV файлу. Якщо не вказано, 
                     буде створено файл з тим самим ім'ям, але з розширенням .csv
    
    Returns:
        Шлях до створеного CSV файлу
    
    Raises:
        ValueError: Якщо формат файлу не підтримується
    """
    ext = os.path.splitext(input_path)[-1].lower()

    if input_path.startswith('https://docs.google'):
        output_path = 'google.csv'

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.csv'
    
    logger.info(f"Конвертація файлу {input_path} у формат CSV: {output_path}")

    try:
        if ext in ['.xlsx', '.xls']:
            _convert_excel_to_csv(input_path, output_path, ext)
        elif ext == '.docx':
            _convert_docx_to_csv(input_path, output_path)
        elif ext == '.doc':
            # convert_doc_to_csv(input_path, output_path)
            raise ValueError(f"Формат {ext} не підтримується!")
        elif ext == '.pdf':
            _convert_pdf_to_csv(input_path, output_path)
        elif ext == '.png' or ext == '.jpg' or ext == '.jpeg':
            _convert_image_to_csv(input_path, output_path)
        elif input_path.startswith('https://docs.google'):
            _convert_google_to_csv(input_path, output_path)
        else:
            raise ValueError(f"Формат {ext} не підтримується!")
        
        logger.info(f"Конвертація завершена успішно. Створено файл: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Помилка при конвертації файлу {input_path}: {str(e)}")
        raise

def _convert_excel_to_csv(input_path, output_path, ext):
    """
    Конвертує Excel файли (.xlsx, .xls) у CSV формат з використанням різних методів.
    Спробує кілька підходів, якщо один з них не спрацює.
    
    Args:
        input_path: Шлях до вхідного Excel файлу
        output_path: Шлях до вихідного CSV файлу
        ext: Розширення файлу (.xlsx або .xls)
    """
    methods = [
        _convert_excel_pandas,
        _convert_excel_openpyxl if ext == '.xlsx' else _convert_excel_xlrd,
        _convert_excel_manual
    ]
    
    last_error = None
    for method in methods:
        try:
            logger.info(f"Спроба конвертації Excel файлу методом {method.__name__}")
            method(input_path, output_path)
            logger.info(f"Конвертація успішна методом {method.__name__}")
            return
        except Exception as e:
            logger.warning(f"Метод {method.__name__} не спрацював: {str(e)}")
            last_error = e
    
    # Якщо жоден метод не спрацював
    raise Exception(f"Не вдалося конвертувати Excel файл. Остання помилка: {str(last_error)}")

def _convert_excel_pandas(input_path, output_path):
    """Конвертація Excel файлу за допомогою pandas"""
    df = pd.read_excel(input_path, engine='openpyxl' if input_path.endswith('.xlsx') else 'xlrd')
    df.to_csv(output_path, index=False, encoding='utf-8')

def _convert_excel_openpyxl(input_path, output_path):
    """Конвертація XLSX файлу за допомогою openpyxl"""
    wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
    sheet = wb.active
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for row in sheet.rows:
            writer.writerow([cell.value for cell in row])

def _convert_excel_xlrd(input_path, output_path):
    """Конвертація XLS файлу за допомогою xlrd"""
    book = xlrd.open_workbook(input_path)
    sheet = book.sheet_by_index(0)
    
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        for row_idx in range(sheet.nrows):
            writer.writerow([sheet.cell_value(row_idx, col_idx) for col_idx in range(sheet.ncols)])

def _convert_excel_manual(input_path, output_path):
    """
    Ручна конвертація Excel файлу шляхом розпакування XLSX як ZIP архіву.
    Цей метод використовується як запасний варіант, якщо інші методи не спрацювали.
    """
    # Створюємо тимчасову копію файлу
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(input_path)[1]) as tmp:
        shutil.copy2(input_path, tmp.name)
        tmp_path = tmp.name
    
    try:
        # Спроба розпакувати XLSX як ZIP архів
        with ZipFile(tmp_path) as zf:
            # Виводимо вміст архіву для діагностики
            file_list = zf.namelist()
            logger.info(f"Вміст архіву: {', '.join(file_list[:10])}...")
            
            # Шукаємо файли worksheet
            worksheet_files = [f for f in file_list if f.startswith('xl/worksheets/sheet')]
            if not worksheet_files:
                raise ValueError("Не знайдено файлів worksheet в архіві")
            
            # Беремо перший лист
            worksheet_file = worksheet_files[0]
            
            # Витягуємо дані з листа
            import xml.etree.ElementTree as ET
            
            # Читаємо XML файл worksheet
            ws_xml = zf.read(worksheet_file)
            ws_root = ET.fromstring(ws_xml)
            
            # Знаходимо всі рядки
            rows = ws_root.findall('.//{*}row')
            logger.info(f"Знайдено {len(rows)} рядків у worksheet")
            
            # Словник для зберігання shared strings, якщо вони є
            shared_strings = []
            
            # Спробуємо знайти shared strings (з урахуванням регістру)
            shared_strings_path = None
            for path in file_list:
                if path.lower() == 'xl/sharedstrings.xml' or path.lower() == 'xl/sharedstrings.xml':
                    shared_strings_path = path
                    break
            
            if shared_strings_path:
                try:
                    logger.info(f"Читання shared strings з {shared_strings_path}")
                    ss_xml = zf.read(shared_strings_path)
                    ss_root = ET.fromstring(ss_xml)
                    
                    # Знаходимо всі елементи si/t
                    for si in ss_root.findall('.//{*}si'):
                        texts = []
                        for t in si.findall('.//{*}t'):
                            if t.text:
                                texts.append(t.text)
                        shared_strings.append(''.join(texts))
                    
                    logger.info(f"Прочитано {len(shared_strings)} shared strings")
                except Exception as e:
                    logger.warning(f"Помилка при читанні shared strings: {str(e)}")
            
            # Відкриваємо CSV файл для запису
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Обробляємо кожен рядок
                for row in rows:
                    csv_row = []
                    
                    # Знаходимо всі комірки в рядку
                    cells = row.findall('.//{*}c')
                    
                    for cell in cells:
                        # Отримуємо значення комірки
                        value = None
                        v_element = cell.find('.//{*}v')
                        
                        if v_element is not None and v_element.text:
                            # Перевіряємо тип комірки
                            cell_type = cell.get('t')
                            
                            if cell_type == 's' and shared_strings:
                                # Це індекс у shared strings
                                try:
                                    idx = int(v_element.text)
                                    if idx < len(shared_strings):
                                        value = shared_strings[idx]
                                    else:
                                        value = v_element.text
                                except (ValueError, IndexError):
                                    value = v_element.text
                            else:
                                # Звичайне значення
                                value = v_element.text
                        
                        csv_row.append(value if value is not None else "")
                    
                    # Записуємо рядок у CSV
                    writer.writerow(csv_row)
                
                logger.info(f"Записано {len(rows)} рядків у CSV файл")
    except Exception as e:
        logger.error(f"Помилка при ручній конвертації Excel файлу: {str(e)}")
        # Створюємо базовий CSV, щоб не повернути порожній файл
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Помилка при конвертації"])
            writer.writerow([str(e)])
    finally:
        # Видаляємо тимчасовий файл
        try:
            os.unlink(tmp_path)
        except Exception as e:
            logger.warning(f"Не вдалося видалити тимчасовий файл: {str(e)}")

def _convert_docx_to_csv(input_path, output_path):
    """Конвертує DOCX файл у CSV формат"""
    doc = Document(input_path)
    
    # Спочатку перевіряємо наявність таблиць
    if doc.tables:
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for table in doc.tables:
                for row in table.rows:
                    writer.writerow([cell.text.strip() for cell in row.cells])
    else:
        # Якщо таблиць немає, записуємо текст параграфів
        text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Text"])
            for line in text.split('\n'):
                if line.strip():
                    writer.writerow([line.strip()])


def _convert_pdf_to_csv(input_path, output_path):
    all_rows = []
    current_brand = None
    pattern_brand = re.compile(r'.*(виробник|гарантія|[A-ZА-Я]{3,}.*\))')

    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            lines = page.extract_text().split('\n')
            for line in lines:
                line = line.strip()

                # Якщо це бренд
                if pattern_brand.match(line):
                    current_brand = line
                    continue

                # Якщо це потенційно рядок з даними акумулятора
                if re.search(r'\d', line) and len(line.split()) > 4:
                    row = re.split(r'\s{2,}|\t', line)
                    all_rows.append([current_brand] + row)

    # Створення DataFrame
    df = pd.DataFrame(all_rows)

    # Додай свій список колонок або збережи без заголовків
    df.to_csv(output_path, index=False, header=False)


def _convert_image_to_csv(input_path, output_path):
   ...

def _convert_google_to_csv(input_path, output_path):
    try:   
        # Витягуємо ID документа та gid з посилання
        doc_id_match = re.search(r'/d/([a-zA-Z0-9-_]+)', input_path)
        if not doc_id_match:
            raise ValueError("Не вдалося витягти ID документа з посилання")
        
        doc_id = doc_id_match.group(1)
        
        # Витягуємо gid (якщо є)
        gid_match = re.search(r'gid=(\d+)', input_path)
        gid = gid_match.group(1) if gid_match else '0'
        
        # Формуємо посилання для експорту CSV
        export_url = f"https://docs.google.com/spreadsheets/d/{doc_id}/export?format=csv&gid={gid}"
        
        # Завантажуємо CSV
        response = requests.get(export_url)
        if response.status_code != 200:
            raise ValueError(f"Помилка при завантаженні CSV: {response.status_code}")
        
        # Зберігаємо CSV у вихідний файл
        with open(output_path, 'wb') as f:
            f.write(response.content)
        
        logger.info(f"Google Spreadsheet успішно конвертовано у CSV: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Помилка при конвертації Google Spreadsheet у CSV: {str(e)}")
        raise